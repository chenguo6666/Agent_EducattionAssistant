import io
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.document_chunk import DocumentChunk
from app.models.uploaded_document import UploadedDocument
from app.rag.ingestion_service import RagIngestionService
from app.rag.text import chunk_text, tokenize_text
from app.schemas.document import DocumentListResponse, DocumentSummaryResponse
from app.services.session_service import SessionService


class DocumentService:
    allowed_suffixes = {".txt", ".md", ".pdf", ".docx"}
    max_file_size = 5 * 1024 * 1024

    def __init__(self) -> None:
        self.session_service = SessionService()
        self.rag_ingestion_service = RagIngestionService()

    async def upload(self, db: Session, user_id: int, file: UploadFile, session_id: str | None = None) -> DocumentSummaryResponse:
        suffix = Path(file.filename or "").suffix.lower()
        if suffix not in self.allowed_suffixes:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Only txt, md, pdf, and docx are supported")

        content = await file.read()
        if not content:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")
        if len(content) > self.max_file_size:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="File is too large; keep it within 5MB")

        session = self.session_service.get_or_create_session(
            db=db,
            user_id=user_id,
            session_id=session_id,
            title_seed=file.filename or "Document Upload",
            title_prefix="资料：",
        )

        extracted_text = self._extract_text(file_name=file.filename or "document", suffix=suffix, content=content)
        if not extracted_text.strip():
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No usable text could be extracted from the file")

        uploads_dir = Path(settings.uploads_dir).resolve()
        uploads_dir.mkdir(parents=True, exist_ok=True)
        document_id = str(uuid.uuid4())
        target_path = uploads_dir / f"{document_id}{suffix}"
        target_path.write_bytes(content)

        record = UploadedDocument(
            id=document_id,
            user_id=user_id,
            session_id=session.id,
            file_name=file.filename or f"document{suffix}",
            file_type=suffix.lstrip("."),
            content_type=file.content_type or "application/octet-stream",
            file_size=len(content),
            file_path=str(target_path),
            extracted_text=extracted_text,
            extraction_status="completed",
        )
        db.add(record)
        db.flush()

        chunk_rows: list[DocumentChunk] = []
        for index, current_chunk in enumerate(chunk_text(extracted_text)):
            chunk_row = DocumentChunk(
                document_id=record.id,
                session_id=session.id,
                user_id=user_id,
                chunk_index=index,
                content=current_chunk,
                keywords_json=tokenize_text(current_chunk),
                embedding_json=None,
            )
            db.add(chunk_row)
            chunk_rows.append(chunk_row)

        db.flush()
        self.rag_ingestion_service.index_chunks(db=db, chunks=chunk_rows)

        self.session_service.touch_session(db=db, session=session)
        db.commit()

        return self._serialize_document(record)

    def list_session_documents(self, db: Session, user_id: int, session_id: str) -> DocumentListResponse:
        session = self.session_service.get_owned_session(db=db, user_id=user_id, session_id=session_id)
        records = (
            db.query(UploadedDocument)
            .filter(UploadedDocument.session_id == session.id, UploadedDocument.user_id == user_id)
            .order_by(UploadedDocument.created_at.asc())
            .all()
        )
        return DocumentListResponse(
            sessionId=session.id,
            documents=[self._serialize_document(record) for record in records],
        )

    def get_session_material_text(self, db: Session, user_id: int, session_id: str | None) -> tuple[str, list[DocumentSummaryResponse]]:
        if not session_id:
            return "", []

        records = (
            db.query(UploadedDocument)
            .filter(UploadedDocument.session_id == session_id, UploadedDocument.user_id == user_id)
            .order_by(UploadedDocument.created_at.asc())
            .all()
        )
        material_parts = [record.extracted_text.strip() for record in records if record.extracted_text.strip()]
        return "\n\n".join(material_parts), [self._serialize_document(record) for record in records]

    def _serialize_document(self, record: UploadedDocument) -> DocumentSummaryResponse:
        snippet = " ".join(record.extracted_text.split())
        if len(snippet) > 80:
            snippet = f"{snippet[:80]}..."

        return DocumentSummaryResponse(
            documentId=record.id,
            sessionId=record.session_id,
            fileName=record.file_name,
            fileType=record.file_type,
            fileSize=record.file_size,
            extractionStatus=record.extraction_status,
            snippet=snippet,
            createdAt=record.created_at,
        )

    def _extract_text(self, file_name: str, suffix: str, content: bytes) -> str:
        if suffix in {".txt", ".md"}:
            return self._decode_text(content)
        if suffix == ".pdf":
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".docx":
            document = DocxDocument(io.BytesIO(content))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Unsupported file format: {file_name}")

    def _decode_text(self, content: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("latin-1", errors="ignore")
